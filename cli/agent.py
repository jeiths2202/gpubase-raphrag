#!/usr/bin/env python3
"""
KMS AI Agent CLI Client - Enterprise Edition

A professional command-line interface for interacting with the KMS AI Agent system.
Features rich terminal UI with cross-platform support.

Usage:
    python -m cli.agent                    # Interactive mode
    python -m cli.agent -q "your question" # Single query mode
    python -m cli.agent --agent ims        # Use specific agent
"""

import argparse
import json
import os
import sys
from typing import Optional, Generator, Dict
from pathlib import Path

try:
    import httpx
except ImportError:
    print("Error: httpx is required. Install with: pip install httpx")
    sys.exit(1)

# Add project root to path
PROJECT_ROOT = Path(__file__).parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from cli.ui import get_ui, EnterpriseUI
from cli.auth import AuthManager
from cli.config import Config
from cli.i18n import get_i18n, I18n


class AgentClient:
    """CLI client for KMS AI Agent API"""

    # Valid agent types from backend (no "auto" - omit agent_type for auto-selection)
    AGENT_TYPES = ["auto", "ims", "rag", "vision", "code", "planner"]
    VALID_API_TYPES = ["ims", "rag", "vision", "code", "planner"]

    # Supported file extensions for attachment
    SUPPORTED_EXTENSIONS = {".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml", ".xml", ".csv", ".log", ".sql", ".sh", ".bat", ".html", ".css", ".pdf", ".docx"}

    # File size limit (500KB for text, 2MB for PDF/DOCX)
    MAX_TEXT_SIZE = 500 * 1024  # 500KB
    MAX_BINARY_SIZE = 2 * 1024 * 1024  # 2MB

    def __init__(self, config: Config, auth: AuthManager, ui: EnterpriseUI):
        self.config = config
        self.auth = auth
        self.ui = ui
        self.session_id: Optional[str] = None
        self.current_agent: str = "auto"
        self.attached_files: Dict[str, str] = {}  # filename -> content

    def attach_file(self, file_path: str) -> tuple[bool, str]:
        """Attach a file for RAG context (text, PDF, DOCX)"""
        path = Path(file_path)

        if not path.exists():
            return False, f"File not found: {file_path}"

        if not path.is_file():
            return False, f"Not a file: {file_path}"

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {ext}. Supported: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"

        try:
            # Check file size first
            file_size = path.stat().st_size
            if ext in {".pdf", ".docx"}:
                if file_size > self.MAX_BINARY_SIZE:
                    return False, f"File too large: {file_size:,} bytes (max 2MB for PDF/DOCX)"
                content = self._extract_document_text(path, ext)
            else:
                if file_size > self.MAX_TEXT_SIZE:
                    return False, f"File too large: {file_size:,} bytes (max 500KB)"
                content = self._read_file_with_encoding(path)

            # Check extracted content size
            if len(content) > self.MAX_TEXT_SIZE:
                return False, f"Extracted content too large: {len(content):,} chars (max 500KB)"

            self.attached_files[path.name] = content
            return True, f"Attached: {path.name} ({len(content):,} chars)"
        except Exception as e:
            return False, f"Failed to read file: {e}"

    def _read_file_with_encoding(self, path: Path) -> str:
        """Read file with encoding detection"""
        encodings = ['utf-8', 'utf-8-sig', 'cp949', 'euc-kr', 'shift_jis', 'latin-1']
        for encoding in encodings:
            try:
                with open(path, 'r', encoding=encoding) as f:
                    return f.read()
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError("Could not decode file with any supported encoding")

    def _extract_document_text(self, path: Path, ext: str) -> str:
        """Extract text from PDF or DOCX files"""
        if ext == ".pdf":
            return self._extract_pdf_text(path)
        elif ext == ".docx":
            return self._extract_docx_text(path)
        else:
            raise ValueError(f"Unsupported document type: {ext}")

    def _extract_pdf_text(self, path: Path) -> str:
        """Extract text from PDF using pypdf"""
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                raise ImportError("PDF support requires 'pypdf' or 'PyPDF2'. Install with: pip install pypdf")

        text_parts = []
        reader = PdfReader(str(path))
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {i + 1}]\n{page_text}")

        if not text_parts:
            raise ValueError("Could not extract text from PDF (may be image-based)")

        return "\n\n".join(text_parts)

    def _extract_docx_text(self, path: Path) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("DOCX support requires 'python-docx'. Install with: pip install python-docx")

        doc = Document(str(path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)

        if not text_parts:
            raise ValueError("Could not extract text from DOCX")

        return "\n".join(text_parts)

    def detach_file(self, filename: str) -> tuple[bool, str]:
        """Detach a file"""
        if filename == "*":
            count = len(self.attached_files)
            self.attached_files.clear()
            return True, f"Detached all {count} files"

        if filename in self.attached_files:
            del self.attached_files[filename]
            return True, f"Detached: {filename}"
        return False, f"File not attached: {filename}"

    def get_attached_files_context(self) -> Optional[str]:
        """Get combined context from attached files"""
        if not self.attached_files:
            return None

        context_parts = []
        for filename, content in self.attached_files.items():
            context_parts.append(f"=== File: {filename} ===\n{content}\n")
        return "\n".join(context_parts)

    def stream_query(self, query: str, agent_type: Optional[str] = None) -> Generator[dict, None, None]:
        """Send query to agent and stream response"""
        agent = agent_type or self.current_agent
        url = f"{self.config.api_url}/agents/stream"

        headers = self.auth.get_headers()
        if not headers:
            self.ui.print_error("Not authenticated. Please login first.")
            return

        payload = {
            "task": query,
            "language": self.config.language
        }

        # Only include agent_type if it's a valid API type (not "auto")
        if agent in self.VALID_API_TYPES:
            payload["agent_type"] = agent

        if self.session_id:
            payload["session_id"] = self.session_id

        # Include attached file context for RAG priority
        file_context = self.get_attached_files_context()
        if file_context:
            payload["file_context"] = file_context

        try:
            with httpx.Client(timeout=float(self.config.timeout)) as client:
                with client.stream("POST", url, json=payload, headers=headers) as response:
                    if response.status_code == 401:
                        self.ui.print_error("Session expired. Please login again.")
                        self.auth.clear_token()
                        return

                    if response.status_code != 200:
                        self.ui.print_error(f"API error: {response.status_code}")
                        return

                    buffer = ""
                    for chunk in response.iter_text():
                        buffer += chunk
                        while "\n" in buffer:
                            line, buffer = buffer.split("\n", 1)
                            if line.startswith("data: "):
                                try:
                                    data = json.loads(line[6:])
                                    yield data
                                except json.JSONDecodeError:
                                    continue

        except httpx.ConnectError:
            self.ui.print_error(f"Cannot connect to server at {self.config.api_url}")
        except httpx.TimeoutException:
            self.ui.print_error("Request timed out")
        except Exception as e:
            self.ui.print_error(f"Error: {e}")

    def query(self, query: str, agent_type: Optional[str] = None, i18n=None) -> str:
        """Send query and collect full response

        Args:
            query: User query
            agent_type: Optional agent type override
            i18n: Optional i18n instance for localized messages
        """
        full_response = ""
        current_type = None
        metadata = {}

        # Show attached files indicator if any
        if self.attached_files:
            file_count = len(self.attached_files)
            file_names = ", ".join(self.attached_files.keys())
            attached_msg = f"📎 Using {file_count} attached file(s): {file_names}"
            if i18n:
                attached_msg = i18n("using_attached_files", count=file_count, files=file_names)
            self.ui.print_info(attached_msg)

        self.ui.start_thinking("Analyzing your request...")

        for chunk in self.stream_query(query, agent_type):
            # API uses "chunk_type" field
            chunk_type = chunk.get("chunk_type", "")
            content = chunk.get("content", "") or ""

            if chunk_type == "thinking":
                if current_type != "thinking":
                    self.ui.update_thinking("Thinking...")
                    current_type = "thinking"

            elif chunk_type == "tool_call":
                tool_name = chunk.get("tool_name", "tool")
                tool_input = chunk.get("tool_input", "")
                self.ui.print_tool_call(tool_name, str(tool_input)[:50] if tool_input else "")
                current_type = "tool_call"

            elif chunk_type == "tool_result":
                self.ui.print_tool_result(len(content))
                self.ui.start_thinking("Processing results...")
                current_type = "tool_result"

            elif chunk_type in ("content", "text"):
                if current_type != "content":
                    self.ui.start_response()
                    current_type = "content"
                # Accumulate response (no immediate output)
                self.ui.stream_response(content)
                full_response += content

            elif chunk_type == "done":
                # Get localized formatting message
                formatting_msg = "Formatting output data..."
                if i18n:
                    formatting_msg = i18n("formatting_output")
                self.ui.end_response(formatting_message=formatting_msg)
                # Extract metadata
                metadata = chunk.get("metadata", {}) or {}
                if metadata.get("session_id"):
                    self.session_id = metadata["session_id"]

            elif chunk_type == "error":
                self.ui.stop_thinking()
                self.ui.print_error(content)

        # Print execution info if available
        if metadata:
            steps = metadata.get("steps", 0)
            exec_time = metadata.get("execution_time", 0) * 1000  # to ms
            if self.session_id and steps:
                self.ui.print_session_info(self.session_id, steps, exec_time)

        return full_response

    def set_agent(self, agent_type: str) -> bool:
        """Set current agent type"""
        if agent_type.lower() in self.AGENT_TYPES:
            old_agent = self.current_agent
            self.current_agent = agent_type.lower()
            self.ui.print_agent_switch(old_agent, self.current_agent)
            return True
        else:
            self.ui.print_error(f"Invalid agent type. Available: {', '.join(self.AGENT_TYPES)}")
            return False

    def new_session(self):
        """Start a new session"""
        self.session_id = None
        self.ui.print_success("New session started")


class CLI:
    """Interactive CLI for Agent"""

    COMMANDS = {
        "/help": "Show this help message",
        "/agent <type>": "Switch agent (auto, ims, rag, vision, code, planner)",
        "/llm": "Show available LLMs and current model",
        "/attach <file>": "Attach a text file for RAG context",
        "/files": "List attached files",
        "/detach <file|*>": "Detach a file (* for all)",
        "/ims-login": "Login to IMS system",
        "/ims-logout": "Logout from IMS system",
        "/new": "Start new session",
        "/status": "Show current status",
        "/clear": "Clear screen",
        "/exit": "Exit CLI",
    }

    def __init__(self, config: Config, ui: EnterpriseUI):
        self.config = config
        self.ui = ui
        self.i18n = get_i18n(config.language)
        self.auth = AuthManager(config)
        self.client = AgentClient(config, self.auth, ui)
        self.running = False

    def show_help(self):
        """Display help information"""
        self.ui.print_help(
            self.COMMANDS,
            AgentClient.AGENT_TYPES,
            self.client.current_agent
        )

    def show_status(self):
        """Display current status"""
        self.ui.print_status(
            server=self.config.api_url,
            agent=self.client.current_agent,
            language=self.config.language,
            logged_in=self.auth.is_authenticated(),
            user=self.auth.get_user_display(),
            session_id=self.client.session_id,
            ims_connected=self.auth.is_ims_authenticated()
        )

    def show_llm_status(self):
        """Display available LLMs and current model"""
        llm_data = self.auth.get_llm_status()
        self.ui.print_llm_status(llm_data, self.client.current_agent)

    def show_attached_files(self):
        """Display list of attached files"""
        self.ui.print_attached_files(self.client.attached_files)

    def check_ims_session(self) -> bool:
        """Check IMS session and auto-login if valid"""
        self.ui.print_info(self.i18n("ims_session_checking"))
        is_valid, _ = self.auth.check_ims_credentials()
        if is_valid:
            self.ui.print_success(self.i18n("ims_session_valid"))
            return True
        return False

    def handle_ims_login(self):
        """Handle IMS login command"""
        self.ui.print_info(self.i18n("ims_login_prompt"))

        try:
            # Get IMS credentials from user (URL is fixed in crawler)
            username = self.ui.get_input(self.i18n("ims_username"))
            password = self.ui.get_password(self.i18n("ims_password"))

            self.ui.print_info(self.i18n("ims_validating"))
            success, message = self.auth.ims_login(username, password)

            if success:
                self.ui.print_success(self.i18n("ims_login_success"))
            else:
                self.ui.print_error(self.i18n("ims_login_failed", error=message))

        except (EOFError, KeyboardInterrupt):
            self.ui.print_warning("Cancelled")

    def handle_ims_logout(self):
        """Handle IMS logout command"""
        success, message = self.auth.ims_logout()
        if success:
            self.ui.print_success(self.i18n("ims_logout_success"))
        else:
            self.ui.print_error(self.i18n("ims_logout_failed"))

    def handle_command(self, cmd: str) -> bool:
        """Handle CLI command. Returns True if should continue."""
        parts = cmd.strip().split(maxsplit=1)
        command = parts[0].lower()
        args = parts[1] if len(parts) > 1 else ""

        if command in ("/exit", "/quit", "/q"):
            self.ui.print_goodbye()
            return False

        elif command in ("/help", "/?", "/h"):
            self.show_help()

        elif command == "/agent":
            if args:
                old_agent = self.client.current_agent
                if self.client.set_agent(args):
                    # When switching to IMS agent, check session
                    if args.lower() == "ims" and old_agent != "ims":
                        if not self.auth.is_ims_authenticated():
                            # Check if there's a valid session
                            if not self.check_ims_session():
                                self.ui.print_warning(self.i18n("ims_not_configured"))
            else:
                self.ui.print_info(f"Current agent: {self.client.current_agent}")
                self.ui.print_info(f"Available: {', '.join(AgentClient.AGENT_TYPES)}")

        elif command == "/llm":
            self.show_llm_status()

        elif command == "/attach":
            if args:
                success, message = self.client.attach_file(args.strip())
                if success:
                    self.ui.print_success(message)
                else:
                    self.ui.print_error(message)
            else:
                self.ui.print_error("Usage: /attach <file_path>")

        elif command == "/files":
            self.show_attached_files()

        elif command == "/detach":
            if args:
                success, message = self.client.detach_file(args.strip())
                if success:
                    self.ui.print_success(message)
                else:
                    self.ui.print_error(message)
            else:
                self.ui.print_error("Usage: /detach <filename|*>")

        elif command == "/ims-login":
            self.handle_ims_login()

        elif command == "/ims-logout":
            self.handle_ims_logout()

        elif command == "/new":
            self.client.new_session()

        elif command == "/status":
            self.show_status()

        elif command == "/clear":
            self.ui.clear()
            self.ui.print_banner()

        else:
            self.ui.print_warning(self.i18n("unknown_command", cmd=command))

        return True

    def login(self, interactive: bool = True) -> bool:
        """Handle login flow"""
        if self.auth.is_authenticated():
            if self.auth.refresh_token():
                self.ui.print_success("Session restored")
                return True

        # Check environment variables first
        env_user = os.environ.get("KMS_USERNAME")
        env_pass = os.environ.get("KMS_PASSWORD")
        if env_user and env_pass:
            self.ui.print_info("Using credentials from environment...")
            if self.auth.login(env_user, env_pass):
                self.ui.print_success(f"Logged in as {env_user}")
                return True

        self.ui.print_login_panel()

        # Try default admin credentials first for development
        if self.config.is_dev_mode():
            self.ui.print_info("Development mode: trying default credentials...")
            if self.auth.login("admin", "SecureAdm1nP@ss2024!"):
                self.ui.print_success("Logged in as admin")
                return True

        # Non-interactive mode - cannot prompt for credentials
        if not interactive:
            self.ui.print_error("Authentication required. Set KMS_USERNAME and KMS_PASSWORD environment variables.")
            return False

        # Manual login
        for attempt in range(3):
            try:
                username = self.ui.get_username()
                password = self.ui.get_password()

                if self.auth.login(username, password):
                    self.ui.print_success(f"Logged in as {username}")
                    return True
                else:
                    remaining = 2 - attempt
                    if remaining > 0:
                        self.ui.print_warning(f"Login failed. {remaining} attempts remaining.")
            except (EOFError, KeyboardInterrupt):
                return False

        self.ui.print_error("Login failed. Please check your credentials.")
        return False

    def run_interactive(self, initial_agent: str = "auto"):
        """Run interactive REPL mode"""
        self.ui.print_banner()

        if not self.login():
            return

        # Set initial agent if specified
        if initial_agent != "auto":
            self.client.set_agent(initial_agent)

        # Check IMS session if using IMS agent
        if self.client.current_agent == "ims":
            if not self.auth.is_ims_authenticated():
                if not self.check_ims_session():
                    self.ui.print_warning(self.i18n("ims_not_configured"))

        self.show_status()
        self.ui.print_info(self.i18n("help_message") + "\n")

        self.running = True
        while self.running:
            try:
                user_input = self.ui.get_prompt(self.client.current_agent)

                if not user_input:
                    continue

                if user_input.startswith("/"):
                    self.running = self.handle_command(user_input)
                else:
                    self.client.query(user_input, i18n=self.i18n)

            except KeyboardInterrupt:
                self.ui.print_info("\nUse /exit to quit or press Ctrl+C again")
                try:
                    input()
                except KeyboardInterrupt:
                    self.ui.print_goodbye()
                    break

            except EOFError:
                self.ui.print_goodbye()
                break

    def run_single_query(self, query: str, agent_type: Optional[str] = None):
        """Run single query mode (non-interactive)"""
        if not self.login(interactive=False):
            return

        if agent_type and agent_type != "auto":
            self.client.set_agent(agent_type)

        # Check IMS session if using IMS agent
        if self.client.current_agent == "ims":
            if not self.auth.is_ims_authenticated():
                if not self.check_ims_session():
                    self.ui.print_warning(self.i18n("ims_login_required"))
                    return

        self.client.query(query, i18n=self.i18n)


def main():
    parser = argparse.ArgumentParser(
        description="KMS AI Agent CLI - Enterprise Edition",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python -m cli.agent                         # Interactive mode
  python -m cli.agent -q "search DFSRRC00"    # Single query
  python -m cli.agent --agent ims -q "list"   # IMS agent query
  python -m cli.agent --server http://host:9000  # Custom server
        """
    )

    parser.add_argument(
        "-q", "--query",
        help="Single query to execute (non-interactive mode)"
    )
    parser.add_argument(
        "-a", "--agent",
        choices=AgentClient.AGENT_TYPES,
        default="auto",
        help="Agent type to use (default: auto)"
    )
    parser.add_argument(
        "-s", "--server",
        default=os.environ.get("KMS_API_URL", "http://localhost:9000"),
        help="API server URL (default: http://localhost:9000)"
    )
    parser.add_argument(
        "-l", "--language",
        choices=["ko", "en", "ja"],
        default="ko",
        help="Response language (default: ko)"
    )
    parser.add_argument(
        "--no-color",
        action="store_true",
        help="Disable colored output"
    )

    args = parser.parse_args()

    # Initialize config
    config = Config(
        api_url=args.server,
        language=args.language,
        use_color=not args.no_color
    )

    # Initialize UI
    ui = get_ui(use_color=not args.no_color)

    # Create CLI and run
    cli = CLI(config, ui)

    if args.query:
        cli.run_single_query(args.query, args.agent)
    else:
        cli.run_interactive(initial_agent=args.agent)


if __name__ == "__main__":
    main()
