"""
Agent API Router
Provides endpoints for agent execution.
"""
from typing import Optional, List
import logging
import json
import io
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from ..core.deps import get_current_user
from ..agents import (
    AgentOrchestrator,
    get_orchestrator,
    AgentRequest,
    AgentResponse,
    AgentType,
    get_tool_registry,
    get_agent_registry,
)

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/agents", tags=["agents"])


# Dependency for orchestrator
async def get_agent_orchestrator() -> AgentOrchestrator:
    """Get the agent orchestrator"""
    return get_orchestrator()


# Response models
class AgentInfo(BaseModel):
    """Agent information"""
    type: str
    name: str
    description: str
    tools: List[str]


class ToolInfo(BaseModel):
    """Tool information"""
    name: str
    description: str


class AgentListResponse(BaseModel):
    """Response for listing agents"""
    agents: List[AgentInfo]


class ToolListResponse(BaseModel):
    """Response for listing tools"""
    tools: List[ToolInfo]


@router.post("/execute", response_model=AgentResponse)
async def execute_agent(
    request: AgentRequest,
    current_user: dict = Depends(get_current_user),
    orchestrator: AgentOrchestrator = Depends(get_agent_orchestrator)
) -> AgentResponse:
    """
    Execute an agent task.

    The agent will be automatically selected based on the task,
    or you can specify a specific agent_type.

    Args:
        request: Agent execution request
        current_user: Authenticated user
        orchestrator: Agent orchestrator

    Returns:
        AgentResponse with answer and metadata
    """
    try:
        user_id = current_user.get("id") or current_user.get("sub")

        response = await orchestrator.execute(request, user_id=user_id)

        return response

    except Exception as e:
        logger.error(f"Agent execution failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"code": "AGENT_ERROR", "message": str(e)}
        )


@router.post("/stream")
async def stream_agent(
    request: AgentRequest,
    current_user: dict = Depends(get_current_user),
    orchestrator: AgentOrchestrator = Depends(get_agent_orchestrator)
):
    """
    Stream agent execution using Server-Sent Events (SSE).

    Provides real-time updates as the agent thinks and acts.

    Args:
        request: Agent execution request
        current_user: Authenticated user
        orchestrator: Agent orchestrator

    Returns:
        StreamingResponse with SSE events
    """
    user_id = current_user.get("id") or current_user.get("sub")

    async def generate():
        try:
            async for chunk in orchestrator.stream(request, user_id=user_id):
                data = chunk.model_dump()
                yield f"data: {json.dumps(data, ensure_ascii=False)}\n\n"
        except Exception as e:
            logger.error(f"Agent streaming failed: {e}")
            yield f"data: {json.dumps({'chunk_type': 'error', 'content': str(e)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


@router.get("/types", response_model=AgentListResponse)
async def list_agents(
    orchestrator: AgentOrchestrator = Depends(get_agent_orchestrator)
) -> AgentListResponse:
    """
    List available agent types.

    Returns information about all registered agents.
    """
    agents = orchestrator.get_available_agents()
    return AgentListResponse(
        agents=[AgentInfo(**a) for a in agents]
    )


@router.get("/tools", response_model=ToolListResponse)
async def list_tools() -> ToolListResponse:
    """
    List available tools.

    Returns information about all registered tools.
    """
    registry = get_tool_registry()
    tools = [
        ToolInfo(name=t.name, description=t.description)
        for t in registry.get_all()
    ]
    return ToolListResponse(tools=tools)


@router.post("/classify")
async def classify_task(
    task: str,
    use_llm: bool = False,
    orchestrator: AgentOrchestrator = Depends(get_agent_orchestrator)
) -> dict:
    """
    Classify a task to determine the appropriate agent type.

    Args:
        task: The task to classify
        use_llm: Whether to use LLM for classification (slower but more accurate)
        orchestrator: Agent orchestrator

    Returns:
        Classification result with agent type and confidence
    """
    if use_llm:
        agent_type = await orchestrator.classify_with_llm(task)
    else:
        agent_type = await orchestrator.classify_task(task)

    return {
        "task": task,
        "agent_type": agent_type.value,
        "method": "llm" if use_llm else "keyword"
    }


# Health check endpoint
@router.get("/health")
async def agent_health():
    """
    Check agent system health.

    Returns status of agent and tool registries.
    """
    try:
        tool_registry = get_tool_registry()
        agent_registry = get_agent_registry()

        return {
            "status": "healthy",
            "agents_registered": len(agent_registry.get_all_types()),
            "tools_registered": len(tool_registry.get_names())
        }
    except Exception as e:
        return {
            "status": "unhealthy",
            "error": str(e)
        }


class ExtractTextResponse(BaseModel):
    """Response for file text extraction"""
    filename: str
    content: str
    size: int
    file_type: str


@router.post("/extract-text", response_model=ExtractTextResponse)
async def extract_text_from_file(
    file: UploadFile = File(..., description="PDF or DOCX file to extract text from"),
    current_user: dict = Depends(get_current_user)
) -> ExtractTextResponse:
    """
    Extract text from PDF or DOCX files for RAG context.

    This endpoint is used by the WebUI to support PDF/DOCX file attachments.
    The extracted text can be used as file_context in agent requests.

    Args:
        file: Uploaded PDF or DOCX file (max 2MB)
        current_user: Authenticated user

    Returns:
        ExtractTextResponse with extracted text content
    """
    # Check file extension
    filename = file.filename or "unknown"
    ext = filename.lower().split(".")[-1] if "." in filename else ""

    if ext not in ("pdf", "docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "INVALID_FILE_TYPE", "message": f"Only PDF and DOCX files are supported. Got: .{ext}"}
        )

    # Read file content
    try:
        content_bytes = await file.read()
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "FILE_READ_ERROR", "message": str(e)}
        )

    # Check file size (2MB limit)
    max_size = 2 * 1024 * 1024
    if len(content_bytes) > max_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "FILE_TOO_LARGE", "message": f"File too large: {len(content_bytes):,} bytes (max 2MB)"}
        )

    # Extract text based on file type
    try:
        if ext == "pdf":
            text_content = _extract_pdf_text(content_bytes)
        else:  # docx
            text_content = _extract_docx_text(content_bytes)
    except ImportError as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail={"code": "MISSING_DEPENDENCY", "message": str(e)}
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "EXTRACTION_ERROR", "message": f"Failed to extract text: {e}"}
        )

    # Check extracted text size (500KB limit)
    max_text_size = 500 * 1024
    if len(text_content) > max_text_size:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail={"code": "CONTENT_TOO_LARGE", "message": f"Extracted content too large: {len(text_content):,} chars (max 500KB)"}
        )

    return ExtractTextResponse(
        filename=filename,
        content=text_content,
        size=len(text_content),
        file_type=ext
    )


def _extract_pdf_text(content_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf"""
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PDF support requires 'pypdf' or 'PyPDF2'. Install with: pip install pypdf")

    text_parts = []
    reader = PdfReader(io.BytesIO(content_bytes))

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(f"[Page {i + 1}]\n{page_text}")

    if not text_parts:
        raise ValueError("Could not extract text from PDF (may be image-based)")

    return "\n\n".join(text_parts)


def _extract_docx_text(content_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx"""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("DOCX support requires 'python-docx'. Install with: pip install python-docx")

    doc = Document(io.BytesIO(content_bytes))
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    if not text_parts:
        raise ValueError("Could not extract text from DOCX")

    return "\n".join(text_parts)
