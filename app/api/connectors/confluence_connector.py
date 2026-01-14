"""
Confluence Connector
Connects to Atlassian Confluence API for pages and spaces.
Supports OAuth 2.0 (3LO) authentication.
"""
import hashlib
import os
import aiohttp
from datetime import datetime
from typing import Optional, List, Dict, Any, AsyncGenerator
from urllib.parse import urlencode

from .base import BaseConnector, ConnectorDocument, ConnectorResult, ConnectorStatus


class ConfluenceConnector(BaseConnector):
    """
    Connector for Atlassian Confluence integration.
    Uses Confluence Cloud REST API with OAuth 2.0 (3LO) authentication.

    Environment Variables:
        ATLASSIAN_CLIENT_ID: OAuth client ID from Atlassian Developer Console
        ATLASSIAN_CLIENT_SECRET: OAuth client secret from Atlassian Developer Console
    """

    # Atlassian OAuth 2.0 endpoints
    OAUTH_AUTH_URL = "https://auth.atlassian.com/authorize"
    OAUTH_TOKEN_URL = "https://auth.atlassian.com/oauth/token"
    ACCESSIBLE_RESOURCES_URL = "https://api.atlassian.com/oauth/token/accessible-resources"

    @property
    def resource_type(self) -> str:
        return "confluence"

    @property
    def display_name(self) -> str:
        return "Confluence"

    @property
    def oauth_scopes(self) -> List[str]:
        """
        OAuth scopes for Confluence API v2 (granular scopes).
        See: https://developer.atlassian.com/cloud/confluence/scopes-for-oauth-2-3LO-and-forge-apps/
        """
        return [
            # Granular scopes for API v2
            "read:page:confluence",           # Read pages
            "read:space:confluence",          # Read spaces
            "read:content-details:confluence", # Read content details
            "read:content.metadata:confluence", # Read content metadata
            "read:attachment:confluence",     # Read attachments
            # Classic scopes for backward compatibility with some endpoints
            "read:confluence-content.all",
            "read:confluence-space.summary",
            # Required for refresh token
            "offline_access"
        ]

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        # OAuth credentials from config or environment
        self.CLIENT_ID = self.config.get("client_id") or os.environ.get("ATLASSIAN_CLIENT_ID", "")
        self.CLIENT_SECRET = self.config.get("client_secret") or os.environ.get("ATLASSIAN_CLIENT_SECRET", "")
        # Cloud ID for API calls (set after OAuth)
        self.cloud_id = self.config.get("cloud_id", "")
        # Base URL from site_url (set after OAuth)
        self.base_url = self.config.get("site_url", "")

        print(f"[ConfluenceConnector] Initialized with cloud_id={self.cloud_id}, base_url={self.base_url}")

    def _get_api_base(self) -> str:
        """Get Confluence API v2 base URL using cloud_id"""
        if self.cloud_id:
            return f"https://api.atlassian.com/ex/confluence/{self.cloud_id}/wiki/api/v2"
        return ""

    def _get_api_base_v1(self) -> str:
        """Get Confluence API v1 base URL (for some endpoints still on v1)"""
        if self.cloud_id:
            return f"https://api.atlassian.com/ex/confluence/{self.cloud_id}/rest/api"
        return ""

    def _get_wiki_base(self) -> str:
        """Get Confluence wiki base URL for web links"""
        if self.base_url:
            return self.base_url
        return ""

    def _get_headers(self) -> Dict[str, str]:
        """Get API request headers with OAuth Bearer token"""
        if self.access_token:
            return {
                "Authorization": f"Bearer {self.access_token}",
                "Accept": "application/json",
                "Content-Type": "application/json"
            }
        return {"Accept": "application/json"}

    # ================== OAuth 2.0 Authentication ==================

    def get_oauth_url(self, redirect_uri: str, state: str) -> str:
        """Generate Atlassian OAuth 2.0 authorization URL"""
        params = {
            "audience": "api.atlassian.com",
            "client_id": self.CLIENT_ID,
            "scope": " ".join(self.oauth_scopes),
            "redirect_uri": redirect_uri,
            "state": state,
            "response_type": "code",
            "prompt": "consent"
        }
        return f"{self.OAUTH_AUTH_URL}?{urlencode(params)}"

    async def exchange_code(self, code: str, redirect_uri: str) -> ConnectorResult:
        """Exchange authorization code for access and refresh tokens"""
        try:
            async with aiohttp.ClientSession() as session:
                # Exchange code for tokens
                async with session.post(
                    self.OAUTH_TOKEN_URL,
                    headers={"Content-Type": "application/json"},
                    json={
                        "grant_type": "authorization_code",
                        "client_id": self.CLIENT_ID,
                        "client_secret": self.CLIENT_SECRET,
                        "code": code,
                        "redirect_uri": redirect_uri
                    }
                ) as resp:
                    if resp.status != 200:
                        error = await resp.text()
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=f"Token exchange failed: {error}"
                        )

                    token_data = await resp.json()
                    access_token = token_data.get("access_token")
                    refresh_token = token_data.get("refresh_token")

                # Get accessible resources (cloud IDs)
                async with session.get(
                    self.ACCESSIBLE_RESOURCES_URL,
                    headers={"Authorization": f"Bearer {access_token}"}
                ) as resp:
                    if resp.status != 200:
                        error = await resp.text()
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=f"Failed to get accessible resources: {error}"
                        )

                    resources = await resp.json()

                    # Find Confluence resource
                    confluence_resource = None
                    for resource in resources:
                        if "confluence" in resource.get("scopes", []) or True:  # Atlassian returns all resources
                            confluence_resource = resource
                            break

                    if not confluence_resource:
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error="No Confluence access found for this account"
                        )

                    cloud_id = confluence_resource.get("id")
                    site_url = confluence_resource.get("url", "")
                    site_name = confluence_resource.get("name", "")

                return ConnectorResult(
                    status=ConnectorStatus.SUCCESS,
                    data={
                        "access_token": access_token,
                        "refresh_token": refresh_token,
                        "cloud_id": cloud_id,
                        "site_url": site_url,
                        "site_name": site_name
                    }
                )

        except Exception as e:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error=str(e)
            )

    async def refresh_access_token(self) -> ConnectorResult:
        """Refresh the access token using refresh token"""
        if not self.refresh_token:
            return ConnectorResult(
                status=ConnectorStatus.AUTH_EXPIRED,
                error="No refresh token available"
            )

        try:
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    self.OAUTH_TOKEN_URL,
                    headers={"Content-Type": "application/json"},
                    json={
                        "grant_type": "refresh_token",
                        "client_id": self.CLIENT_ID,
                        "client_secret": self.CLIENT_SECRET,
                        "refresh_token": self.refresh_token
                    }
                ) as resp:
                    if resp.status != 200:
                        return ConnectorResult(
                            status=ConnectorStatus.AUTH_EXPIRED,
                            error="Token refresh failed"
                        )

                    token_data = await resp.json()
                    return ConnectorResult(
                        status=ConnectorStatus.SUCCESS,
                        data={
                            "access_token": token_data.get("access_token"),
                            "refresh_token": token_data.get("refresh_token", self.refresh_token)
                        }
                    )

        except Exception as e:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error=str(e)
            )

    async def validate_connection(self) -> ConnectorResult:
        """Validate Confluence OAuth connection using API v2"""
        if not self.access_token:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error="No access token available"
            )

        if not self.cloud_id:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error="No cloud_id configured"
            )

        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(
                    f"{self._get_api_base()}/spaces",
                    headers=self._get_headers(),
                    params={"limit": 1}
                ) as resp:
                    if resp.status == 200:
                        return ConnectorResult(
                            status=ConnectorStatus.SUCCESS,
                            message="Connection validated"
                        )
                    elif resp.status == 401:
                        return ConnectorResult(
                            status=ConnectorStatus.AUTH_EXPIRED,
                            error="Access token expired"
                        )
                    else:
                        error = await resp.text()
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=error
                        )
        except Exception as e:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error=str(e)
            )

    # ================== Document Operations ==================

    async def list_documents(
        self,
        path: Optional[str] = None,
        modified_since: Optional[datetime] = None
    ) -> AsyncGenerator[Dict[str, Any], None]:
        """List pages and attachments from Confluence spaces"""
        print(f"[ConfluenceConnector] list_documents called: path={path}, modified_since={modified_since}")
        print(f"[ConfluenceConnector] Using cloud_id={self.cloud_id}, base_url={self.base_url}")
        print(f"[ConfluenceConnector] Access token present: {bool(self.access_token)}")

        doc_count = 0
        try:
            async with aiohttp.ClientSession() as session:
                # Get spaces
                spaces = await self._list_spaces(session)
                print(f"[ConfluenceConnector] Found {len(spaces)} spaces")

                for space in spaces:
                    space_id = space["id"]
                    space_key = space["key"]
                    space_name = space["name"]
                    print(f"[ConfluenceConnector] Processing space: {space_key} ({space_name}), id={space_id}")

                    # Filter by path (space key) if specified
                    if path and space_key != path:
                        continue

                    # Get pages in space
                    async for page in self._list_pages(
                        session, space_id, space_key, modified_since
                    ):
                        page_id = page["id"]
                        page_title = page["title"]

                        # Yield page itself
                        doc_count += 1
                        yield {
                            "external_id": page_id,
                            "title": page_title,
                            "path": f"{space_name}",
                            "space_key": space_key,
                            "url": f"{self.base_url}{page['_links']['webui']}",
                            "type": page["type"],
                            "status": page.get("status"),
                            "version": page.get("version", {}).get("number"),
                            "modified_time": page.get("version", {}).get("when")
                        }

                        # Also list attachments for this page
                        async for attachment in self._list_attachments(
                            session, page_id, page_title, space_name
                        ):
                            doc_count += 1
                            yield {
                                "external_id": f"att_{attachment['id']}",
                                "title": attachment["title"],
                                "path": f"{space_name} / {page_title}",
                                "space_key": space_key,
                                "url": f"{self.base_url}{attachment['download_link']}",
                                "type": "attachment",
                                "mime_type": attachment["media_type"],
                                "parent_page_id": page_id,
                                "version": attachment.get("version"),
                                "modified_time": attachment.get("modified_time")
                            }

            print(f"[ConfluenceConnector] list_documents completed: total {doc_count} documents")

        except Exception as e:
            print(f"[ConfluenceConnector] List documents error: {e}")
            import traceback
            traceback.print_exc()

    async def _list_spaces(
        self,
        session: aiohttp.ClientSession
    ) -> List[Dict]:
        """List accessible spaces using Confluence API v2"""
        spaces = []

        api_base = self._get_api_base()
        if not api_base:
            print(f"[ConfluenceConnector] ERROR: No API base URL (cloud_id={self.cloud_id})")
            return spaces

        print(f"[ConfluenceConnector] Listing spaces from {api_base}/spaces")

        try:
            cursor = None
            limit = 50

            while True:
                url = f"{api_base}/spaces"
                params = {
                    "limit": limit,
                    "status": "current"
                }
                if cursor:
                    params["cursor"] = cursor

                print(f"[ConfluenceConnector] GET {url} params={params}")

                async with session.get(
                    url,
                    headers=self._get_headers(),
                    params=params
                ) as resp:
                    if resp.status != 200:
                        error_text = await resp.text()
                        print(f"[ConfluenceConnector] List spaces failed: status={resp.status}, error={error_text}")
                        break

                    data = await resp.json()
                    results = data.get("results", [])
                    print(f"[ConfluenceConnector] Got {len(results)} spaces")

                    # Convert v2 format to expected format
                    for space in results:
                        spaces.append({
                            "id": space.get("id"),
                            "key": space.get("key"),
                            "name": space.get("name"),
                            "type": space.get("type"),
                            "status": space.get("status")
                        })

                    # Check for next page using cursor-based pagination
                    next_link = data.get("_links", {}).get("next")
                    if not next_link or len(results) < limit:
                        break

                    # Extract cursor from next link
                    import re
                    cursor_match = re.search(r'cursor=([^&]+)', next_link)
                    if cursor_match:
                        cursor = cursor_match.group(1)
                    else:
                        break

        except Exception as e:
            print(f"[ConfluenceConnector] List spaces error: {e}")
            import traceback
            traceback.print_exc()

        return spaces

    async def _list_pages(
        self,
        session: aiohttp.ClientSession,
        space_id: str,
        space_key: str,
        modified_since: Optional[datetime] = None
    ) -> AsyncGenerator[Dict, None]:
        """List pages in a space using Confluence API v2"""
        print(f"[ConfluenceConnector] Listing pages in space: {space_key} (id={space_id})")
        try:
            cursor = None
            limit = 50
            page_count = 0

            while True:
                url = f"{self._get_api_base()}/spaces/{space_id}/pages"
                params = {
                    "limit": limit,
                    "status": "current"
                }
                if cursor:
                    params["cursor"] = cursor

                async with session.get(
                    url,
                    headers=self._get_headers(),
                    params=params
                ) as resp:
                    if resp.status != 200:
                        error_text = await resp.text()
                        print(f"[ConfluenceConnector] List pages failed for space {space_key}: status={resp.status}, error={error_text}")
                        break

                    data = await resp.json()
                    results = data.get("results", [])
                    print(f"[ConfluenceConnector] Got {len(results)} pages in space {space_key}")

                    for page in results:
                        # Filter by modification time using v2 format
                        if modified_since:
                            version_date = page.get("version", {}).get("createdAt")
                            if version_date:
                                try:
                                    page_modified = datetime.fromisoformat(
                                        version_date.replace("Z", "+00:00")
                                    )
                                    if page_modified <= modified_since:
                                        continue
                                except Exception:
                                    pass

                        page_count += 1
                        # Convert v2 format to expected format
                        yield {
                            "id": page.get("id"),
                            "title": page.get("title"),
                            "type": "page",
                            "status": page.get("status"),
                            "version": page.get("version", {}),
                            "_links": {
                                "webui": page.get("_links", {}).get("webui", f"/wiki/spaces/{space_key}/pages/{page.get('id')}")
                            }
                        }

                    # Check for next page using cursor-based pagination
                    next_link = data.get("_links", {}).get("next")
                    if not next_link or len(results) < limit:
                        break

                    # Extract cursor from next link
                    import re
                    cursor_match = re.search(r'cursor=([^&]+)', next_link)
                    if cursor_match:
                        cursor = cursor_match.group(1)
                    else:
                        break

            print(f"[ConfluenceConnector] Total pages yielded from space {space_key}: {page_count}")

        except Exception as e:
            print(f"[ConfluenceConnector] List pages error: {e}")
            import traceback
            traceback.print_exc()

    async def _list_attachments(
        self,
        session: aiohttp.ClientSession,
        page_id: str,
        page_title: str,
        space_name: str
    ) -> AsyncGenerator[Dict, None]:
        """List attachments for a page"""
        try:
            start = 0
            limit = 50

            while True:
                async with session.get(
                    f"{self._get_api_base()}/content/{page_id}/child/attachment",
                    headers=self._get_headers(),
                    params={
                        "start": start,
                        "limit": limit,
                        "expand": "version"
                    }
                ) as resp:
                    if resp.status != 200:
                        break

                    data = await resp.json()
                    results = data.get("results", [])

                    for attachment in results:
                        # Only process supported file types
                        media_type = attachment.get("metadata", {}).get("mediaType", "")
                        if media_type in ["application/pdf", "text/plain", "text/markdown",
                                          "application/msword",
                                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                            yield {
                                "id": attachment["id"],
                                "title": attachment["title"],
                                "media_type": media_type,
                                "page_id": page_id,
                                "page_title": page_title,
                                "space_name": space_name,
                                "download_link": attachment.get("_links", {}).get("download", ""),
                                "version": attachment.get("version", {}).get("number"),
                                "modified_time": attachment.get("version", {}).get("when")
                            }

                    if len(results) < limit:
                        break
                    start += limit

        except Exception as e:
            print(f"[ConfluenceConnector] List attachments error: {e}")

    async def fetch_document(self, external_id: str) -> ConnectorResult:
        """Fetch Confluence page or attachment content using API v2"""
        # Check if this is an attachment
        if external_id.startswith("att_"):
            return await self._fetch_attachment(external_id[4:])  # Remove "att_" prefix

        # Fetch regular page using v2 API
        try:
            async with aiohttp.ClientSession() as session:
                # First get page metadata
                async with session.get(
                    f"{self._get_api_base()}/pages/{external_id}",
                    headers=self._get_headers(),
                    params={
                        "body-format": "storage"
                    }
                ) as resp:
                    if resp.status == 404:
                        return ConnectorResult(
                            status=ConnectorStatus.NOT_FOUND,
                            error="Page not found"
                        )
                    if resp.status != 200:
                        error_text = await resp.text()
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=f"API error: {resp.status} - {error_text}"
                        )

                    page = await resp.json()

                # Extract content from body (v2 format)
                body = page.get("body", {}).get("storage", {}).get("value", "")
                content = self._html_to_text(body)
                sections = self._parse_confluence_sections(body)

                # Parse modification time from v2 format
                modified_at = None
                version_date = page.get("version", {}).get("createdAt")
                if version_date:
                    try:
                        modified_at = datetime.fromisoformat(
                            version_date.replace("Z", "+00:00")
                        )
                    except Exception:
                        pass

                # Get space info
                space_id = page.get("spaceId", "")
                space_key = ""
                space_name = ""

                # Fetch space details if we have space_id
                if space_id:
                    try:
                        async with session.get(
                            f"{self._get_api_base()}/spaces/{space_id}",
                            headers=self._get_headers()
                        ) as space_resp:
                            if space_resp.status == 200:
                                space_data = await space_resp.json()
                                space_key = space_data.get("key", "")
                                space_name = space_data.get("name", "")
                    except Exception:
                        pass

                # Build path
                path = space_name or "Confluence"

                content_hash = hashlib.md5(content.encode()).hexdigest()

                # Build web URL
                webui_link = page.get("_links", {}).get("webui", "")
                if webui_link and not webui_link.startswith("http"):
                    webui_link = f"{self.base_url}{webui_link}"

                doc = ConnectorDocument(
                    external_id=external_id,
                    title=page.get("title", "Untitled"),
                    content=content,
                    external_url=webui_link,
                    path=path,
                    mime_type="text/html",
                    modified_at=modified_at,
                    sections=sections,
                    content_hash=content_hash,
                    metadata={
                        "confluence_id": external_id,
                        "space_id": space_id,
                        "space_key": space_key,
                        "space_name": space_name,
                        "version": page.get("version", {}).get("number"),
                        "type": "page",
                        "author": page.get("version", {}).get("authorId")
                    }
                )

                return ConnectorResult(
                    status=ConnectorStatus.SUCCESS,
                    data=doc
                )

        except Exception as e:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error=str(e)
            )

    async def _fetch_attachment(self, attachment_id: str) -> ConnectorResult:
        """Fetch and parse Confluence attachment content"""
        try:
            async with aiohttp.ClientSession() as session:
                # First get attachment metadata
                async with session.get(
                    f"{self._get_api_base()}/content/{attachment_id}",
                    headers=self._get_headers(),
                    params={"expand": "version,container,space"}
                ) as resp:
                    if resp.status == 404:
                        return ConnectorResult(
                            status=ConnectorStatus.NOT_FOUND,
                            error="Attachment not found"
                        )
                    if resp.status != 200:
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=f"API error: {resp.status}"
                        )

                    attachment = await resp.json()

                title = attachment.get("title", "Untitled")
                download_link = attachment.get("_links", {}).get("download", "")
                media_type = attachment.get("metadata", {}).get("mediaType", "")
                container = attachment.get("container", {})
                space = attachment.get("space", {})

                # Parse modification time
                modified_at = None
                version_when = attachment.get("version", {}).get("when")
                if version_when:
                    modified_at = datetime.fromisoformat(
                        version_when.replace("Z", "+00:00")
                    )

                # Download attachment content
                download_url = f"{self.base_url}{download_link}"
                async with session.get(
                    download_url,
                    headers=self._get_headers()
                ) as download_resp:
                    if download_resp.status != 200:
                        return ConnectorResult(
                            status=ConnectorStatus.ERROR,
                            error=f"Failed to download attachment: {download_resp.status}"
                        )

                    file_content = await download_resp.read()

                # Extract text based on media type
                content = ""
                if media_type == "application/pdf":
                    content = self._extract_pdf_text(file_content)
                elif media_type in ["text/plain", "text/markdown"]:
                    content = file_content.decode("utf-8", errors="ignore")
                elif media_type in ["application/msword",
                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                    content = self._extract_docx_text(file_content)
                else:
                    content = f"[Binary attachment: {title}]"

                content_hash = hashlib.md5(content.encode()).hexdigest()
                path = f"{space.get('name', '')} / {container.get('title', '')}"

                doc = ConnectorDocument(
                    external_id=f"att_{attachment_id}",
                    title=title,
                    content=content,
                    external_url=download_url,
                    path=path,
                    mime_type=media_type,
                    modified_at=modified_at,
                    sections=[{"title": title, "content": content, "type": "document"}],
                    content_hash=content_hash,
                    metadata={
                        "confluence_attachment_id": attachment_id,
                        "parent_page_id": container.get("id"),
                        "parent_page_title": container.get("title"),
                        "space_key": space.get("key"),
                        "space_name": space.get("name"),
                        "media_type": media_type,
                        "type": "attachment"
                    }
                )

                return ConnectorResult(
                    status=ConnectorStatus.SUCCESS,
                    data=doc
                )

        except Exception as e:
            return ConnectorResult(
                status=ConnectorStatus.ERROR,
                error=str(e)
            )

    def _extract_pdf_text(self, pdf_content: bytes) -> str:
        """Extract text from PDF content"""
        try:
            import io
            from pypdf import PdfReader

            pdf_file = io.BytesIO(pdf_content)
            reader = PdfReader(pdf_file)

            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return "\n\n".join(text_parts)

        except Exception as e:
            print(f"[ConfluenceConnector] PDF extraction error: {e}")
            return f"[PDF extraction failed: {e}]"

    def _extract_docx_text(self, docx_content: bytes) -> str:
        """Extract text from DOCX content"""
        try:
            import io
            from docx import Document

            docx_file = io.BytesIO(docx_content)
            doc = Document(docx_file)

            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            return "\n\n".join(text_parts)

        except Exception as e:
            print(f"[ConfluenceConnector] DOCX extraction error: {e}")
            return f"[DOCX extraction failed: {e}]"

    def _parse_confluence_sections(self, html: str) -> List[Dict]:
        """Parse Confluence storage format HTML into sections"""
        import re

        sections = []
        current_section = {"title": "", "content": [], "type": "text"}

        # Handle Confluence-specific elements
        # Extract structured content

        # Find headings
        heading_pattern = r'<h([1-6])[^>]*>(.*?)</h\1>'
        parts = re.split(heading_pattern, html)

        i = 0
        while i < len(parts):
            if i + 2 < len(parts) and parts[i+1].isdigit():
                # Found a heading
                if current_section["content"]:
                    current_section["content"] = self._html_to_text(
                        "\n".join(current_section["content"])
                    )
                    sections.append(current_section)

                level = int(parts[i+1])
                title = self._html_to_text(parts[i+2])
                current_section = {
                    "title": title.strip(),
                    "content": [],
                    "type": "text",
                    "level": level
                }
                i += 3
            else:
                # Regular content
                content = parts[i]
                if content.strip():
                    # Check for code blocks
                    code_blocks = re.findall(
                        r'<ac:structured-macro ac:name="code"[^>]*>.*?<ac:plain-text-body><!\[CDATA\[(.*?)\]\]></ac:plain-text-body>.*?</ac:structured-macro>',
                        content, re.DOTALL
                    )
                    if code_blocks:
                        for code in code_blocks:
                            sections.append({
                                "title": "",
                                "content": code,
                                "type": "code"
                            })
                        # Remove code blocks from content
                        content = re.sub(
                            r'<ac:structured-macro ac:name="code"[^>]*>.*?</ac:structured-macro>',
                            '', content, flags=re.DOTALL
                        )

                    # Check for tables
                    tables = re.findall(r'<table[^>]*>.*?</table>', content, re.DOTALL)
                    if tables:
                        for table in tables:
                            table_content = self._parse_confluence_table(table)
                            sections.append({
                                "title": "",
                                "content": table_content,
                                "type": "table"
                            })
                        content = re.sub(r'<table[^>]*>.*?</table>', '', content, flags=re.DOTALL)

                    if content.strip():
                        current_section["content"].append(content)
                i += 1

        if current_section["content"]:
            current_section["content"] = self._html_to_text(
                "\n".join(current_section["content"])
            )
            sections.append(current_section)

        return sections

    def _parse_confluence_table(self, table_html: str) -> List[List[str]]:
        """Parse HTML table into rows"""
        import re

        rows = []
        row_matches = re.findall(r'<tr[^>]*>(.*?)</tr>', table_html, re.DOTALL)

        for row in row_matches:
            cells = []
            cell_matches = re.findall(r'<t[hd][^>]*>(.*?)</t[hd]>', row, re.DOTALL)
            for cell in cell_matches:
                cell_text = self._html_to_text(cell)
                cells.append(cell_text.strip())
            if cells:
                rows.append(cells)

        return rows
